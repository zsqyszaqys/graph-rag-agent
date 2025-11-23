from typing import Dict, List, Tuple, Optional
from docx import Document
import codecs
import os
import PyPDF2
import csv
import json
import yaml
from yaml import CLoader as Loader

class FileReader:
    """
    文件读取器，支持多种文件格式：
        - TXT (文本文件)
        - PDF (PDF文档)
        - MD (Markdown文件)
        - DOCX (Word文档)
        - DOC (旧版Word文档)
        - CSV (CSV文件)
        - JSON (JSON文件)
        - YAML/YML (YAML文件)
    处理流程:
            FileReader 初始化
            ↓
        调用 read_files(extensions, recursive)
            ↓
            ├─ recursive=True ─→ _read_files_recursive() ─┐
            │                                              │
            └─ recursive=False → _process_files_in_dir()─┘
                                                           ↓
                                            根据文件扩展名选择读取方法
                                                           ↓
                                ┌──────────────────────────┼──────────────────────┐
                                ↓                          ↓                      ↓
                           _read_txt()              _read_pdf()            _read_docx()
                           _read_csv()              _read_json()           _read_yaml()
                                ↓                          ↓                      ↓
                                └──────────────────────────┴──────────────────────┘
                                                           ↓
                                    返回 List[Tuple[文件名, 文本内容]]
    """

    def __init__(self, directory_path:str):
        """
        初始化文件读取器
        :param directory_path:文件目录路径
        """
        self.directory_path = directory_path

    def read_files(self, file_extensions: Optional[List[str]] = None, recursive: bool = True) -> List[Tuple[str, str]]:
        """
         读取指定扩展名的文件
        :param file_extensions:文件扩展名列表，如 ['.txt', '.pdf']，如不指定则读取所有支持的格式
        :param recursive:是否递归读取子目录，默认为True
        :return:List[Tuple[str, str]]: 文件名和内容的元组列表
        """
        supported_extensions = {
            '.txt': self._read_txt,
            '.pdf': self._read_pdf,
            '.md': self._read_markdown,
            '.docx': self._read_docx,
            '.doc': self._read_doc,
            '.csv': self._read_csv,
            '.json': self._read_json,
            '.yaml': self._read_yaml,
            '.yml': self._read_yaml,
        }

        # 如未指定扩展名，则使用所有支持的扩展名
        if file_extensions is None:
            file_extensions = list(supported_extensions.keys())

        results = []
        try:
            if recursive:
                # 递归读取所有文件
                results = self._read_files_recursive(self.directory_path, file_extensions, supported_extensions)
                print(f"递归读取目录完成，总共读取了 {len(results)} 个文件")
            else:
                # 仅读取当前目录的文件
                all_filenames = os.listdir(self.directory_path)
                print(f"当前目录中共有 {len(all_filenames)} 个文件")

                results = self._process_files_in_dir(self.directory_path, all_filenames, file_extensions,
                                                     supported_extensions)
                print(f"总共读取了 {len(results)} 个文件")
        except Exception as e:
            print(f"列出目录 {self.directory_path} 中的文件时出错: {str(e)}")

        return results

    def _read_files_recursive(self, root_dir:str, file_extensions:List[str], supported_extensions:Dict)->List[Tuple[str, str]]:
        """
        递归读取目录及其子目录中的文件
        :param root_dir:当前处理的目录路径
        :param file_extensions:要处理的文件扩展名列表
        :param supported_extensions:支持的文件扩展名及对应处理函数
        :return:List[Tuple[str, str]]: 文件名和内容的元组列表
        """
        results = []

        try:
            # 遍历目录内容
            for item in os.listdir(root_dir):
                item_path = os.path.join(root_dir, item)

                # 如果是目录，递归处理
                if os.path.isdir(item_path):
                    print(f"递归进入子目录: {item_path}")
                    sub_results = self._read_files_recursive(item_path, file_extensions, supported_extensions)
                    results.extend(sub_results)
                # 如果是文件，处理文件
                elif os.path.isfile(item_path):
                    file_ext = os.path.splitext(item)[1].lower()

                    if file_ext in file_extensions:
                        # 获取相对与根目录的路径
                        rel_path = os.path.relpath(item_path, self.directory_path)

                        print(f"处理文件: {rel_path} (类型: {file_ext})")

                        # 使用对应的读取方法处理文件
                        if file_ext in supported_extensions:
                            try:
                                content = supported_extensions[file_ext](item_path)
                                # 存储相对路径而不是仅文件名，以便区分不同目录中的同名文件
                                results.append((rel_path, content))
                                print(f"成功读取文件: {rel_path}, 内容长度: {len(content)}")
                            except Exception as e:
                                print(f"读取文件 {rel_path} 时出错: {str(e)}")
        except Exception as e:
            print(f"列出目录 {root_dir} 中的文件时出错: {str(e)}")

        return results

    def _process_files_in_dir(self, directory: str, filenames: List[str], file_extensions: List[str],
                              supported_extensions: Dict) -> List[Tuple[str, str]]:
        """
        处理指定目录中的文件（不递归）
        :param directory:目录路径
        :param filenames:文件名列表
        :param file_extensions:要处理的文件扩展名列表
        :param supported_extensions:支持的文件扩展名及对应处理函数
        :return:List[Tuple[str, str]]: 文件名和内容的元组列表
        """

        results = []

        for filename in filenames:
            file_ext = os.path.splitext(filename)[1].lower()

            if file_ext in file_extensions:
                file_path = os.path.join(directory, filename)
                print(f"处理文件: {filename} (类型: {file_ext})")

                # 使用对应的读取方法处理文件
                if file_ext in supported_extensions:
                    try:
                        content = supported_extensions[file_ext](file_path)
                        results.append((filename, content))
                        print(f"成功读取文件: {filename}, 内容长度: {len(content)}")
                    except Exception as e:
                        print(f"读取文件 {filename} 时出错: {str(e)}")

        return results

    def _read_txt(self, file_path:str)->str:
        """读取txt文件"""
        try:
            with codecs.open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                content = file.read()
            return content
        except Exception as e:
            print(f"读取TXT文件 {os.path.basename(file_path)} 失败: {str(e)}")

            # 尝试使用其他的编码
            try:
                with open(file_path, 'rb') as f: # 二进制模式读取
                    raw_data = f.read(10240) # 读取前10kb
                    try:
                        import chardet
                        result = chardet.detect(raw_data)
                        encoding = result['encoding'] if result['encoding'] else 'gbk'
                    except:
                        encoding = 'gbk' # 如果chardet不可用，默认使用gbk

                # 用检测到的编码重新读取
                with codecs.open(file_path, encoding=encoding, errors='replace') as file:
                    content = file.read()
                return content
            except Exception as e2:
                print(f"尝试使用其他编码读取失败: {str(e2)}")
                return f"[无法读取文件内容: {str(e)}]"

    def _read_pdf(self, file_path:str)->str:
        """读取pdf文件"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_txt = page.extract_text() or ""
                        text += page_txt + "\n\n"
                    except Exception as e:
                        print(f"读取PDF文件 {os.path.basename(file_path)} 的第 {page_num + 1} 页失败: {str(e)}")
                        text += f"[第 {page_num + 1} 页无法读取]\n\n"

                return text
        except Exception as e:
            print(f"读取PDF文件 {os.path.basename(file_path)} 失败: {str(e)}")
            return f"[无法读取PDF文件内容: {str(e)}]"

    def _read_markdown(self, file_path:str)->str:
        """读取markdown文件"""
        try:
            with codecs.open(file_path, "r", encoding='utf-8', errors='replace') as file:
                md_content = file.read()
                return md_content
        except Exception as e:
            print(f"读取Markdown文件 {os.path.basename(file_path)} 失败: {str(e)}")
            return f"[无法读取Markdown文件内容: {str(e)}]"

    def _read_docx(self, file_path:str)->str:
        """读取Word文档(.docx)"""
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            print(f"读取Word文档(.docx) {os.path.basename(file_path)} 失败: {str(e)}")
            return f"[无法读取Word文档内容: {str(e)}]"

    def _read_doc(self, file_path: str) -> str:
        """
        读取旧版Word文档(.doc)
        首先尝试使用Windows特有的方法，如果失败则使用跨平台的方法
        """
        content = ""

        # 方法1: 尝试使用win32com (仅Windows)
        try:
            import win32com.client

            print(f"尝试使用win32com读取.doc文件: {os.path.basename(file_path)}")
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            doc_abs_path = os.path.abspath(file_path)
            doc = word.Documents.Open(doc_abs_path)
            content = doc.Content.Text
            doc.Close()
            word.Quit()

            if content and content.strip():
                print(f"使用win32com成功读取.doc文件")
                return content
        except ImportError:
            print("win32com不可用，这不是Windows系统")
        except Exception as e:
            print(f"使用win32com读取.doc失败: {str(e)}")

        # 方法2: 尝试使用textract (跨平台)
        try:
            import textract
            print(f"尝试使用textract读取.doc文件: {os.path.basename(file_path)}")
            content = textract.process(file_path).decode('utf-8')

            if content and content.strip():
                print(f"使用textract成功读取.doc文件")
                return content
        except ImportError:
            print("textract不可用，请安装: pip install textract")
        except Exception as e:
            print(f"使用textract读取.doc失败: {str(e)}")

        # 方法3: 尝试使用python-docx (不完全兼容.doc，但有时可以部分读取)
        try:
            from docx import Document
            print(f"尝试使用python-docx读取.doc文件: {os.path.basename(file_path)}")
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = '\n'.join(full_text)

            if content and content.strip():
                print(f"使用python-docx部分读取.doc文件成功")
                return content
        except ImportError:
            print("python-docx不可用，请安装: pip install python-docx")
        except Exception as e:
            print(f"尝试使用python-docx读取.doc失败: {str(e)}")

        # 所有方法都失败，返回警告信息
        warning_msg = f"[警告: 无法读取.doc文件 {os.path.basename(file_path)}，请安装相关依赖或转换为.docx格式]"
        print(warning_msg)
        return warning_msg

    def _read_csv(self, file_path:str)->str:
        """
        读取csv文件并转化为纯文本(无结构化处理)
        """
        try:
            text = []
            with open(file_path, "r", encoding='utf-8', errors="replace") as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    text.append(','.join(row))
            return '\n'.join(text)
        except Exception as e:
            print(f"读取CSV文件 {os.path.basename(file_path)} 失败: {str(e)}")
            # 尝试其他的编码
            try:
                with open(file_path, "rb") as f:
                    try:
                        import chardet
                        raw_data = f.read(10240)
                        result = chardet.detect(raw_data)
                        encoding = result['encoding'] if result['encoding'] else 'gbk'
                    except:
                        encoding = 'gbk'

                text = []
                with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                    csv_reader = csv.reader(file)
                    for row in csv_reader:
                        text.append(','.join(row))
                    return '\n'.join(text)
            except Exception as e2:
                print(f"尝试使用其他编码读取CSV失败: {str(e2)}")
                return f"[无法读取CSV文件内容: {str(e)}]"

    def read_csv_as_dicts(self, file_path:str)->List[Dict]:
        """
        读取csv文件并返回字典列表
        :return:List[Dict]: CSV数据的字典列表，每一行为一个字典
        """
        try:
            results = []
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                csv_reader = csv.DictReader(file)
                for row in csv_reader:
                    results.append(dict(row))
                return results
        except Exception as e:
            print(f"读取CSV文件为字典列表时出错: {str(e)}")
            return []

    def _read_json(self, file_path:str)->str:
        """读取json文件并返回文本格式"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                # 加载为对象然后再转为格式化的字符串
                data = json.load(file)
                return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"读取JSON文件 {os.path.basename(file_path)} 失败: {str(e)}")
            return f"[无法读取JSON文件内容: {str(e)}]"

    def read_json_as_dict(self, file_path: str) -> Dict:
        """
        读取JSON文件并返回字典/列表对象
        :return:Dict/List: JSON数据对象
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return json.load(file)
        except Exception as e:
            print(f"读取JSON文件为字典时出错: {str(e)}")
            return {}

    def _read_yaml(self, file_path: str) -> str:
        """读取YAML文件并返回文本格式"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                data = yaml.load(file, Loader=Loader)
                # 先转为JSON字符串以获得更易读的格式
                return yaml.dump(data, allow_unicode=True, default_flow_style=False)
        except Exception as e:
            print(f"读取YAML文件 {os.path.basename(file_path)} 失败: {str(e)}")
            return f"[无法读取YAML文件内容: {str(e)}]"

    def read_yaml_as_dict(self, file_path: str) -> Dict:
        """
        读取YAML文件并返回字典对象
        :return: Dict: YAML数据对象
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return yaml.load(file, Loader=Loader)
        except Exception as e:
            print(f"读取YAML文件为字典时出错: {str(e)}")
            return {}

    def read_txt_files(self) -> List[Tuple[str, str]]:
        """读取所有txt文件"""
        return self.read_files(['.txt'])

    def list_all_files(self, recursive: bool = True) -> List[str]:
        """
        列出目录中的所有文件
        :param recursive:是否递归列出子目录中的文件，默认为True
        :return: List[str]: 文件路径列表（相对于根目录）
        """
        files = []

        try:
            if recursive:
                # 递归遍历所有子目录
                for root, _, filenames in os.walk(self.directory_path):
                    for filename in filenames:
                        # 获取相对于根目录的路径
                        rel_path = os.path.relpath(os.path.join(root, filename), self.directory_path)
                        files.append(rel_path)
            else:
                # 只列出当前目录下的文件
                files = os.listdir(self.directory_path)
        except Exception as e:
            print(f"列出目录文件时出错: {str(e)}")

        return files
